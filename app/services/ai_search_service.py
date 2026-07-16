"""
AI Search service implementing Retrieval Augmented Generation (RAG) using Groq and HuggingFace local embeddings.
"""

from __future__ import annotations

import os
import time
from typing import Any, Dict, List, Optional, Tuple
from uuid import UUID
from difflib import SequenceMatcher

from sqlalchemy import select
from sqlalchemy.ext.asyncio import AsyncSession

from app.config import get_settings
from app.logging_config import get_logger
from app.models.article import Article, ArticleStatus
from app.models.resource import Resource
from app.models.search_history import SearchHistory

logger = get_logger(__name__)
settings = get_settings()

try:
    from langchain_community.vectorstores import Chroma
    from langchain_community.embeddings.fastembed import FastEmbedEmbeddings
    from langchain_core.documents import Document
    from langchain_text_splitters import RecursiveCharacterTextSplitter
    from langchain_core.prompts import PromptTemplate
    HAS_LANGCHAIN = True
except ImportError as e:
    logger.warning("langchain_imports_failed", error=str(e))
    HAS_LANGCHAIN = False


def extract_text_from_file_path(file_path: str, file_type: str) -> str:
    ext = file_type.lower().lstrip(".")
    if ext in ("txt", "md"):
        with open(file_path, "r", encoding="utf-8", errors="ignore") as f:
            return f.read()
    elif ext == "pdf":
        from pypdf import PdfReader
        reader = PdfReader(file_path)
        pages_text = []
        for page in reader.pages:
            t = page.extract_text()
            if t:
                pages_text.append(t)
        return "\n".join(pages_text)
    elif ext == "docx":
        import docx
        d = docx.Document(file_path)
        paragraphs_text = [p.text for p in d.paragraphs]
        return "\n".join(paragraphs_text)
    return ""


class AISearchService:
    """Handles vector search indexing and RAG answer generation using Groq and HuggingFace Embeddings."""

    def __init__(self, db: AsyncSession):
        self.db = db
        self._embeddings = None
        self._groq_client = None

        # Load local FastEmbed embeddings (bypasses PyTorch DLL issues, requires no API key)
        if HAS_LANGCHAIN:
            try:
                self._embeddings = FastEmbedEmbeddings(
                    model_name="BAAI/bge-small-en-v1.5"
                )
            except Exception as ex:
                logger.error("fastembed_embeddings_init_failed", error=str(ex))

        # Initialize official Groq client
        if settings.GROQ_API_KEY:
            try:
                from groq import Groq
                self._groq_client = Groq(api_key=settings.GROQ_API_KEY)
            except Exception as ex:
                logger.error("groq_client_init_failed", error=str(ex))

    def _get_vector_store(self) -> Chroma:
        """Helper to safely initialize vector store using the dedicated chromadb_groq directory."""
        return Chroma(
            persist_directory=settings.CHROMA_DB_DIR,
            embedding_function=self._embeddings
        )

    async def ingest_knowledge_base(self) -> dict:
        """Fetch all published articles, seed resources, and active knowledge documents, chunk them, embed, and index."""
        if not HAS_LANGCHAIN or not self._embeddings:
            return {"success": False, "message": "Langchain or embeddings models are not initialized."}

        # 1. Fetch data from DB
        articles_stmt = select(Article).where(Article.status == ArticleStatus.PUBLISHED)
        articles_result = await self.db.execute(articles_stmt)
        articles = articles_result.scalars().all()

        resources_stmt = select(Resource)
        resources_result = await self.db.execute(resources_stmt)
        resources = resources_result.scalars().all()

        from app.models.knowledge_document import KnowledgeDocument
        docs_stmt = select(KnowledgeDocument).where(KnowledgeDocument.indexing_status == "active")
        docs_result = await self.db.execute(docs_stmt)
        kb_docs = docs_result.scalars().all()

        # 2. Convert to LangChain Documents
        documents = []
        for art in articles:
            doc = Document(
                page_content=f"Title: {art.title}\nExcerpt: {art.excerpt or ''}\nContent:\n{art.content_markdown}",
                metadata={
                    "source": "article",
                    "id": str(art.id),
                    "title": art.title,
                    "slug": art.slug,
                }
            )
            documents.append(doc)

        for res in resources:
            doc = Document(
                page_content=f"Resource Title: {res.title}\nDescription: {res.description or ''}\nURL: {res.url or ''}\nType: {res.resource_type.value}",
                metadata={
                    "source": "resource",
                    "id": str(res.id),
                    "title": res.title,
                    "url": res.url,
                }
            )
            documents.append(doc)

        for kb_doc in kb_docs:
            try:
                text = extract_text_from_file_path(kb_doc.file_path, kb_doc.file_type)
                if text.strip():
                    doc = Document(
                        page_content=text,
                        metadata={
                            "source": "document",
                            "id": str(kb_doc.id),
                            "title": kb_doc.title,
                        }
                    )
                    documents.append(doc)
            except Exception as e:
                logger.error("failed_to_extract_during_ingest", error=str(e), id=str(kb_doc.id))

        if not documents:
            return {"success": True, "indexed_chunks": 0, "message": "No documents found to index."}

        # 3. Chunking
        text_splitter = RecursiveCharacterTextSplitter(chunk_size=500, chunk_overlap=50)
        chunks = text_splitter.split_documents(documents)

        # 4. Vector DB indexing
        try:
            vector_store = self._get_vector_store()
            try:
                collection_data = vector_store.get()
                if collection_data and collection_data.get('ids'):
                    vector_store.delete(ids=collection_data['ids'])
            except Exception:
                pass
            Chroma.from_documents(
                documents=chunks,
                embedding=self._embeddings,
                persist_directory=settings.CHROMA_DB_DIR
            )
            logger.info("indexing_completed", chunks=len(chunks))
            return {"success": True, "indexed_chunks": len(chunks)}
        except Exception as e:
            logger.error("vector_indexing_failed", error=str(e))
            return {"success": False, "message": f"Vector indexing failed: {str(e)}"}


    async def index_article(self, article: Article):
        """Index or update a single article in the vector store in a background task."""
        if not HAS_LANGCHAIN or not self._embeddings:
            return

        # Extract values to safe types to prevent accessing closed session in background thread
        try:
            status_val = article.status.value if hasattr(article.status, "value") else str(article.status)
        except Exception:
            status_val = str(article.status)

        article_data = {
            "id": str(article.id),
            "status": status_val,
            "title": str(article.title),
            "excerpt": str(article.excerpt or ""),
            "content_markdown": str(article.content_markdown or ""),
            "slug": str(article.slug),
        }

        async def _bg_task():
            try:
                vector_store = self._get_vector_store()
                
                try:
                    vector_store.delete(where={"id": article_data["id"]})
                except Exception:
                    pass

                # Embed only if published
                if article_data["status"] == "published":
                    doc = Document(
                        page_content=f"Title: {article_data['title']}\nExcerpt: {article_data['excerpt']}\nContent:\n{article_data['content_markdown']}",
                        metadata={
                            "source": "article",
                            "id": article_data["id"],
                            "title": article_data["title"],
                            "slug": article_data["slug"],
                        }
                    )
                    text_splitter = RecursiveCharacterTextSplitter(chunk_size=500, chunk_overlap=50)
                    chunks = text_splitter.split_documents([doc])
                    vector_store.add_documents(chunks)
                    logger.info("Article indexed.")
                    logger.info("Embeddings updated.")
            except Exception as e:
                logger.error("article_auto_embed_failed", error=str(e))

        import asyncio
        asyncio.create_task(_bg_task())

    async def index_resource(self, resource: Resource):
        """Index or update a single resource in the vector store in a background task."""
        if not HAS_LANGCHAIN or not self._embeddings:
            return

        # Extract values to safe types to prevent accessing closed session in background thread
        try:
            res_type_val = resource.resource_type.value if hasattr(resource.resource_type, "value") else str(resource.resource_type)
        except Exception:
            res_type_val = str(resource.resource_type)

        resource_data = {
            "id": str(resource.id),
            "title": str(resource.title),
            "description": str(resource.description or ""),
            "url": str(resource.url or ""),
            "resource_type": res_type_val,
        }

        async def _bg_task():
            try:
                vector_store = self._get_vector_store()
                
                try:
                    vector_store.delete(where={"id": resource_data["id"]})
                except Exception:
                    pass

                doc = Document(
                    page_content=f"Resource Title: {resource_data['title']}\nDescription: {resource_data['description']}\nURL: {resource_data['url']}\nType: {resource_data['resource_type']}",
                    metadata={
                        "source": "resource",
                        "id": resource_data["id"],
                        "title": resource_data["title"],
                        "url": resource_data["url"],
                    }
                )
                text_splitter = RecursiveCharacterTextSplitter(chunk_size=500, chunk_overlap=50)
                chunks = text_splitter.split_documents([doc])
                vector_store.add_documents(chunks)
                logger.info("Resource indexed.")
                logger.info("Embeddings updated.")
            except Exception as e:
                logger.error("resource_auto_embed_failed", error=str(e))

        import asyncio
        asyncio.create_task(_bg_task())

    async def index_document(self, doc_id: UUID):
        """Extract text, chunk, embed, and store in ChromaDB for an uploaded KnowledgeDocument."""
        if not HAS_LANGCHAIN or not self._embeddings:
            return

        async def _bg_task():
            from app.database import AsyncSessionLocal
            from app.models.knowledge_document import KnowledgeDocument
            from datetime import datetime, timezone
            
            async with AsyncSessionLocal() as session:
                try:
                    stmt = select(KnowledgeDocument).where(KnowledgeDocument.id == doc_id)
                    result = await session.execute(stmt)
                    db_doc = result.scalar_one_or_none()
                    if not db_doc:
                        logger.error("doc_not_found_for_indexing", id=str(doc_id))
                        return

                    db_doc.indexing_status = "processing"
                    await session.commit()

                    # Extract text using helper
                    try:
                        text = extract_text_from_file_path(db_doc.file_path, db_doc.file_type)
                    except Exception as extract_err:
                        logger.error("text_extraction_failed", error=str(extract_err), id=str(doc_id))
                        db_doc.indexing_status = "error"
                        await session.commit()
                        return

                    if not text.strip():
                        logger.error("no_text_extracted", id=str(doc_id))
                        db_doc.indexing_status = "error"
                        await session.commit()
                        return

                    # Chunking
                    doc = Document(
                        page_content=text,
                        metadata={
                            "source": "document",
                            "id": str(db_doc.id),
                            "title": db_doc.title,
                        }
                    )
                    text_splitter = RecursiveCharacterTextSplitter(chunk_size=500, chunk_overlap=50)
                    chunks = text_splitter.split_documents([doc])

                    # Index in ChromaDB
                    vector_store = self._get_vector_store()
                    try:
                        vector_store.delete(where={"id": str(db_doc.id)})
                    except Exception:
                        pass

                    vector_store.add_documents(chunks)

                    # Update status
                    db_doc.chunks_count = len(chunks)
                    db_doc.indexed_at = datetime.now(timezone.utc)
                    db_doc.indexing_status = "active"
                    await session.commit()
                    logger.info("document_indexing_success", id=str(doc_id), chunks=len(chunks))

                except Exception as e:
                    logger.error("document_indexing_failed", error=str(e), id=str(doc_id))
                    try:
                        stmt = select(KnowledgeDocument).where(KnowledgeDocument.id == doc_id)
                        result = await session.execute(stmt)
                        db_doc = result.scalar_one_or_none()
                        if db_doc:
                            db_doc.indexing_status = "error"
                            await session.commit()
                    except Exception:
                        pass

        import asyncio
        asyncio.create_task(_bg_task())

    async def delete_document(self, item_id: UUID):

        """Delete all chunks matching the given id from the vector store in a background task."""
        if not HAS_LANGCHAIN or not self._embeddings:
            return

        doc_id = str(item_id)

        async def _bg_task():
            try:
                vector_store = self._get_vector_store()
                vector_store.delete(where={"id": doc_id})
                logger.info("Embeddings removed.")
            except Exception as e:
                logger.error("document_deletion_from_vector_failed", error=str(e))

        import asyncio
        asyncio.create_task(_bg_task())

    async def rag_search(
        self, query: str, user_id: Optional[UUID] = None, force_ai: bool = False
    ) -> Dict[str, Any]:
        """Perform semantic search and return generated response with sources using Groq."""
        start_time = time.time()
        # Pre-initialize SearchHistory
        history = None
        if user_id:
            history = SearchHistory(
                user_id=user_id,
                query=query.strip()[:1000],
                results_count=0,
                search_type="ai",
            )
            self.db.add(history)
            await self.db.flush()

        if not HAS_LANGCHAIN or not self._embeddings:
            fallback_answer, fallback_sources = await self._database_fallback_search(query)
            if user_id and history:
                history.results_count = len(fallback_sources)
                history.search_type = "ai_fallback"
                await self.db.flush()
            return {
                "answer": fallback_answer,
                "sources": fallback_sources,
                "engine": "fallback-db-lookup"
            }

        # Check Groq API client
        if not self._groq_client:
            friendly_error = "The AI service is temporarily unavailable. Please try again later."
            if user_id and history:
                history.search_type = "ai_error"
                await self.db.flush()
            return {
                "answer": friendly_error,
                "sources": [],
                "engine": "error-fallback"
            }

        try:
            # Safely get Chroma vector store instance
            vector_store = self._get_vector_store()

            # Lazy loading / Ingestion check
            is_empty = True
            try:
                if vector_store._collection.count() > 0:
                    is_empty = False
            except Exception:
                pass

            if is_empty:
                logger.info("chromadb_empty_ingesting_with_local_embeddings")
                await self.ingest_knowledge_base()

            # Initialize timing variables
            embedding_time_ms = 0
            retrieval_time_ms = 0
            generation_time_ms = 0

            # 1. Explicitly measure embedding time
            t0 = time.time()
            self._embeddings.embed_query(query)
            embedding_time_ms = int((time.time() - t0) * 1000)

            # 2. Measure retrieval time or handle force_ai
            t1 = time.time()
            if force_ai:
                res = {
                    "answer": "",
                    "sources": [],
                    "engine": "general-groq-fallback",
                    "search_type": "ai_general",
                    "confidence": "High",
                    "knowledge_source": "General AI Knowledge"
                }
                
                try:
                    system_prompt = "You are Startup Navigator AI, an expert advisor for startup founders. Answer the user's question using your general knowledge."
                    t2 = time.time()
                    completion = self._groq_client.chat.completions.create(
                        model="llama-3.1-8b-instant",
                        messages=[
                            {"role": "system", "content": system_prompt},
                            {"role": "user", "content": query}
                        ],
                        temperature=0.3,
                        max_tokens=1024
                    )
                    generation_time_ms = int((time.time() - t2) * 1000)
                    res["answer"] = completion.choices[0].message.content
                except Exception as e:
                    logger.error("groq_generation_failed", error=str(e))
                    res["answer"] = "Unable to generate an AI response right now. Please try again later."
                    res["engine"] = "error-fallback"

            else:
                docs_with_scores = vector_store.similarity_search_with_relevance_scores(query, k=4)
                raw_retrieval_time = int((time.time() - t1) * 1000)
                retrieval_time_ms = max(0, raw_retrieval_time - embedding_time_ms)
                
                SIMILARITY_THRESHOLD = 0.65
                relevant_docs_with_scores = [(doc, score) for doc, score in docs_with_scores if score > SIMILARITY_THRESHOLD]
                
                res = None
                # Route 1: Relevant documents found in local Knowledge Base (RAG)
                if relevant_docs_with_scores:
                    # Deduplicate chunks based on text similarity (>80%)
                    unique_chunks = []
                    for doc, score in relevant_docs_with_scores:
                        if len(unique_chunks) >= 3:
                            break
                        
                        is_duplicate = False
                        for existing_doc, _ in unique_chunks:
                            similarity = SequenceMatcher(None, doc.page_content, existing_doc.page_content).ratio()
                            if similarity > 0.8:
                                is_duplicate = True
                                break
                        
                        if not is_duplicate:
                            unique_chunks.append((doc, score))
                            
                    # Build merged context
                    context_text = "\n\n".join([f"Source: {doc.metadata.get('title', 'Document')}\n{doc.page_content}" for doc, _ in unique_chunks])
                    
                    # Deduplicate sources by document id
                    max_score = 0
                    sources = []
                    seen_doc_ids = set()
                    
                    for doc, score in unique_chunks:
                        if score > max_score:
                            max_score = score
                            
                        doc_id = doc.metadata.get("id")
                        if doc_id not in seen_doc_ids:
                            seen_doc_ids.add(doc_id)
                            similarity_pct = int(round(score * 100))
                            sources.append({
                                "id": doc_id,
                                "title": doc.metadata.get("title"),
                                "slug": doc.metadata.get("slug"),
                                "url": doc.metadata.get("url"),
                                "source_type": doc.metadata.get("source"),
                                "similarity": similarity_pct,
                                "section": doc.metadata.get("title")
                            })
                    
                    # Call LLM for Summarization
                    summary_answer = ""
                    try:
                        system_prompt = (
                            "You are Startup Navigator AI, an expert advisor for startup founders. "
                            "Answer the user's question using ONLY the provided context below. "
                            "Format your response with the Answer, and Key Points (using bullets) if appropriate. "
                            "Never repeat the same sentence or paragraph. Provide a clean, concise, summarized answer."
                        )
                        t2 = time.time()
                        completion = self._groq_client.chat.completions.create(
                            model="llama-3.1-8b-instant",
                            messages=[
                                {"role": "system", "content": system_prompt},
                                {"role": "user", "content": f"Context:\n{context_text}\n\nQuestion: {query}"}
                            ],
                            temperature=0.1,
                            max_tokens=1024
                        )
                        generation_time_ms = int((time.time() - t2) * 1000)
                        summary_answer = completion.choices[0].message.content
                    except Exception as e:
                        logger.error("groq_rag_generation_failed", error=str(e))
                        summary_answer = "Found relevant information, but failed to generate a summarized response."
                    
                    # Determine confidence
                    if max_score > 0.6:
                        confidence = "High"
                    elif max_score > 0.4:
                        confidence = "Medium"
                    else:
                        confidence = "Low"

                    res = {
                        "answer": summary_answer,
                        "sources": sources,
                        "engine": "chroma-extract",
                        "search_type": "ai_rag",
                        "confidence": confidence,
                        "knowledge_source": "Knowledge Base"
                    }

                # Route 2: No relevant documents found (DO NOT CALL GROQ)
                if res is None:
                    res = {
                        "answer": "",
                        "sources": [],
                        "engine": "chroma-empty",
                        "search_type": "ai_rag_empty",
                        "confidence": "Low",
                        "knowledge_source": "Knowledge Base",
                        "status": "not_found"
                    }

            response_time_ms = int((time.time() - start_time) * 1000)
            res["response_time_ms"] = response_time_ms
            
            logger.info(
                "rag_search_query_log", 
                query=query, 
                retrieved_chunks_count=len(res.get("sources", [])), 
                knowledge_source=res.get("knowledge_source"), 
                response_time_ms=response_time_ms
            )
            
            if user_id and history:
                history.results_count = len(res.get("sources", []))
                history.search_type = res.get("search_type", "ai_rag")
                history.ai_answer = res.get("answer")
                history.knowledge_source = res.get("knowledge_source")
                history.source_documents = res.get("sources", [])
                history.response_time_ms = response_time_ms
                history.embedding_time_ms = embedding_time_ms
                history.retrieval_time_ms = retrieval_time_ms
                history.generation_time_ms = generation_time_ms
                await self.db.flush()
                
            return res

        except Exception as e:
            logger.error("groq_generation_failed", error=str(e), query=query)
            friendly_error = "The AI service is temporarily unavailable. Please try again later."
            if user_id and history:
                history.search_type = "ai_error"
                await self.db.flush()
            return {
                "answer": friendly_error,
                "sources": [],
                "engine": "error-fallback"
            }

    async def _database_fallback_search(self, query: str) -> Tuple[str, List[dict]]:
        """Fallback keyword search directly against SQL database when LLM/Vector stores are offline."""
        search_term = f"%{query.strip()}%"
        stmt = select(Article).where(
            Article.status == ArticleStatus.PUBLISHED,
            Article.title.ilike(search_term)
        ).limit(3)
        
        result = await self.db.execute(stmt)
        articles = result.scalars().all()
        
        if not articles:
            return (
                "I couldn't find any direct match in my knowledge base. Please try rephrasing your search.",
                []
            )

        titles = ", ".join([a.title for a in articles])
        answer = f"I found some articles that might help you: {titles}. Here is an excerpt from the first match: {articles[0].excerpt or ''}"
        sources = [
            {
                "id": str(a.id),
                "title": a.title,
                "slug": a.slug,
                "source_type": "article"
            }
            for a in articles
        ]
        return answer, sources
